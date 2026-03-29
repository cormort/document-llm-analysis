import json
import os

import structlog
from app.utils.excel_processor import smart_read_excel
from docx import Document
from pypdf import PdfReader

logger = structlog.get_logger()


class DocumentExtractionService:
    """Service for extracting text from PDF, Word, and Excel files."""

    async def extract_text(
        self, file_path: str, start_page: int = 1, end_page: int = None
    ) -> str:
        """Unified entry point for text extraction"""
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".pdf":
                return self._extract_pdf(file_path, start_page, end_page)
            elif ext in [".docx", ".doc"]:
                return self._extract_word(file_path)
            elif ext in [".xlsx", ".xls", ".csv"]:
                return self._extract_excel(file_path)
            elif ext == ".json":
                with open(file_path, encoding="utf-8") as f:
                    data = json.load(f)
                return json.dumps(data, ensure_ascii=False, indent=2)
            elif ext in [".txt", ".md", ".py", ".js", ".html"]:
                with open(file_path, encoding="utf-8") as f:
                    return f.read()
            else:
                return f"[Unsupported file format: {ext}]"
        except Exception as e:
            logger.error(f"Extraction failed for {file_path}", error=str(e))
            return f"[Error reading file: {str(e)}]"

    def _extract_pdf(
        self, file_path: str, start_page: int = 1, end_page: int | None = None
    ) -> str:
        """Extract text and tables from PDF using pdfplumber.

        This method handles mixed content PDFs (text + tables) by:
        1. Detecting tables on each page
        2. Extracting table content as Markdown format
        3. Extracting remaining text content
        """
        try:
            import pdfplumber
        except ImportError:
            # Fallback to pypdf if pdfplumber not available
            logger.warning("pdfplumber not available, falling back to pypdf")
            return self._extract_pdf_fallback(file_path, start_page, end_page)

        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)

                if start_page < 1:
                    start_page = 1
                if end_page is None or end_page > total_pages:
                    end_page = total_pages
                if start_page > end_page:
                    start_page = end_page

                output = (
                    f"=== PDF: {total_pages} 頁 (提取: {start_page}-{end_page}) ===\n"
                )

                for page_num in range(start_page - 1, end_page):
                    page = pdf.pages[page_num]
                    output += f"\n\n[第 {page_num + 1} 頁]\n"

                    # 1. 提取表格
                    tables = page.extract_tables()
                    if tables:
                        for i, table in enumerate(tables):
                            output += f"\n📊 表格 {i + 1}:\n"
                            output += self._table_to_markdown(table)
                            output += "\n"

                    # 2. 提取非表格文字 (使用 filter 移除表格區域)
                    # 取得表格邊界框
                    table_bboxes = [t.bbox for t in page.find_tables()]

                    def not_within_any_table(obj):
                        """Check if object is outside all table bounding boxes."""
                        for bbox in table_bboxes:
                            if (
                                obj["x0"] >= bbox[0]
                                and obj["x1"] <= bbox[2]
                                and obj["top"] >= bbox[1]
                                and obj["bottom"] <= bbox[3]
                            ):
                                return False
                        return True

                    # 過濾掉表格區域的字元
                    filtered_page = page.filter(not_within_any_table)
                    text = filtered_page.extract_text() or ""

                    if text.strip():
                        output += f"\n📝 文字內容:\n{text}\n"

                return output

        except Exception as e:
            logger.error("pdfplumber extraction failed", error=str(e))
            return self._extract_pdf_fallback(file_path, start_page, end_page)

    def _table_to_markdown(self, table: list) -> str:
        """Convert pdfplumber table to Markdown format."""
        if not table or not table[0]:
            return ""

        # Clean cells
        def clean_cell(cell):
            if cell is None:
                return ""
            return str(cell).replace("\n", " ").strip()

        rows = [[clean_cell(cell) for cell in row] for row in table]

        # Build markdown table
        if not rows:
            return ""

        # Header
        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"

        # Body
        body_rows = [
            "| " + " | ".join(row) + " |"
            for row in rows[1:]
            if any(cell for cell in row)
        ]

        return header + "\n" + separator + "\n" + "\n".join(body_rows)

    def _extract_pdf_fallback(
        self, file_path: str, start_page: int = 1, end_page: int | None = None
    ) -> str:
        """Fallback PDF extraction using pypdf (no table support)."""
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)

        if start_page < 1:
            start_page = 1
        if end_page is None or end_page > total_pages:
            end_page = total_pages
        if start_page > end_page:
            start_page = end_page

        text = f"=== PDF Page Range: {start_page} to {end_page} (Total: {total_pages}) ===\n"

        for i in range(start_page - 1, end_page):
            text += f"\n[Page {i + 1}]\n" + reader.pages[i].extract_text() + "\n"
        return text

    def _extract_word(self, file_path: str) -> str:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    [cell.text.strip().replace("\n", " ") for cell in row.cells]
                )
                text += f"\n| {row_text} |"
        return text

    def _extract_excel(self, file_path: str) -> str:
        # Use our smart processor logic
        import openpyxl
        import pandas as pd

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".csv":
            try:
                df = pd.read_csv(file_path)
                return df.to_markdown(index=False)
            except Exception as e:
                logger.error("CSV extraction failed", error=str(e))
                return f"[Error reading CSV: {str(e)}]"

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            all_text = ""
            for sheet_name in wb.sheetnames:
                df = smart_read_excel(
                    file_path, header_rows=1, sheet_name=sheet_name, unmerge=True
                )
                if not df.empty:
                    all_text += f"\n\n=== Sheet: {sheet_name} ===\n"
                    all_text += df.to_markdown(index=False)
            return all_text
        except Exception as e:
            logger.error("Excel extraction failed", error=str(e))
            # Fallback
            try:
                df = smart_read_excel(file_path)
                return df.to_markdown(index=False)
            except Exception as ex:
                logger.error("Fallback Excel read failed", error=str(ex))
                return f"[Error reading Excel: {str(e)}]"


document_service = DocumentExtractionService()
