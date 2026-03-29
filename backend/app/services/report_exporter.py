"""
Report Export Utilities
Generate HTML, Markdown and DOCX reports from AI interpretations
"""

import os
import re
from datetime import datetime

import markdown


def generate_docx_report(title: str, content: str, output_path: str) -> str:
    """Generate a Word (.docx) report from markdown content.
    
    Args:
        title: Report title
        content: Markdown-formatted report content
        output_path: Path where the .docx file should be saved
    
    Returns:
        The output_path of the generated file
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # --- Document styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Noto Sans TC"
    font.size = Pt(11)

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x36, 0x5E, 0xD8)

    # Timestamp
    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_run = ts_para.add_run(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    ts_run.font.size = Pt(9)
    ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # --- Parse and add content line by line ---
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        # Horizontal rule
        elif stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 40)
        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_markdown(para, text)
        # Numbered list
        elif re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_inline_markdown(para, text)
        # Empty line
        elif stripped == "":
            doc.add_paragraph()
        # Normal paragraph
        else:
            para = doc.add_paragraph()
            _add_inline_markdown(para, stripped)

        i += 1

    doc.save(output_path)
    return output_path


def _add_inline_markdown(para, text: str) -> None:
    """Add inline markdown (bold, italic) to a paragraph."""
    # Split on bold (**text**) patterns
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def generate_html_report(title: str, content: str, chart_html: str = None) -> str:
    """Generate a styled HTML report from markdown content."""
    
    # Convert markdown to HTML
    html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
    
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    html_template = f"""
<!DOCTYPE html>
<html lang="zh-TW">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} - AI 分析報告</title>
    <style>
        :root {{
            --primary: #667EEA;
            --primary-dark: #5A67D8;
            --bg-dark: #1a1a2e;
            --bg-card: #16213e;
            --text: #e4e6eb;
            --text-secondary: #b0b3b8;
        }}
        
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Noto Sans TC', sans-serif;
            background: linear-gradient(135deg, var(--bg-dark) 0%, #0f0f23 100%);
            color: var(--text);
            line-height: 1.8;
            min-height: 100vh;
            padding: 40px 20px;
        }}
        
        .container {{
            max-width: 900px;
            margin: 0 auto;
        }}
        
        .header {{
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 30px;
            border-bottom: 1px solid rgba(102, 126, 234, 0.3);
        }}
        
        .header h1 {{
            font-size: 2rem;
            background: linear-gradient(135deg, var(--primary) 0%, #764ba2 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            margin-bottom: 10px;
        }}
        
        .header .timestamp {{
            color: var(--text-secondary);
            font-size: 0.9rem;
        }}
        
        .header .badge {{
            display: inline-block;
            background: rgba(102, 126, 234, 0.2);
            color: var(--primary);
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 0.8rem;
            margin-top: 10px;
        }}
        
        .content {{
            background: var(--bg-card);
            border-radius: 16px;
            padding: 40px;
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3);
            border: 1px solid rgba(102, 126, 234, 0.1);
        }}
        
        .content h1, .content h2, .content h3 {{
            color: var(--primary);
            margin-top: 24px;
            margin-bottom: 16px;
        }}
        
        .content h1 {{ font-size: 1.6rem; }}
        .content h2 {{ font-size: 1.4rem; }}
        .content h3 {{ font-size: 1.2rem; }}
        
        .content p {{
            margin-bottom: 16px;
            color: var(--text);
        }}
        
        .content ul, .content ol {{
            margin: 16px 0;
            padding-left: 24px;
        }}
        
        .content li {{
            margin-bottom: 8px;
        }}
        
        .content strong {{
            color: #fff;
        }}
        
        .content code {{
            background: rgba(102, 126, 234, 0.15);
            color: #a78bfa;
            padding: 2px 6px;
            border-radius: 4px;
            font-family: 'SF Mono', Monaco, monospace;
            font-size: 0.9em;
        }}
        
        .content pre {{
            background: #0d1117;
            border-radius: 8px;
            padding: 16px;
            overflow-x: auto;
            margin: 16px 0;
        }}
        
        .content pre code {{
            background: none;
            padding: 0;
        }}
        
        .content table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}
        
        .content th, .content td {{
            padding: 12px;
            text-align: left;
            border-bottom: 1px solid rgba(255, 255, 255, 0.1);
        }}
        
        .content th {{
            background: rgba(102, 126, 234, 0.2);
            color: var(--primary);
            font-weight: 600;
        }}
        
        .content tr:hover {{
            background: rgba(102, 126, 234, 0.05);
        }}
        
        .chart-container {{
            margin: 30px 0;
            border-radius: 12px;
            overflow: hidden;
        }}
        
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid rgba(102, 126, 234, 0.2);
            color: var(--text-secondary);
            font-size: 0.85rem;
        }}
        
        @media print {{
            body {{
                background: white;
                color: #333;
            }}
            .content {{
                box-shadow: none;
                border: 1px solid #ddd;
            }}
            .header h1 {{
                background: none;
                -webkit-text-fill-color: var(--primary);
                color: var(--primary);
            }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📊 {title}</h1>
            <div class="timestamp">生成時間：{timestamp}</div>
            <span class="badge">🤖 AI 專業分析報告</span>
        </div>
        
        <div class="content">
            {f'<div class="chart-container">{chart_html}</div>' if chart_html else ''}
            {html_content}
        </div>
        
        <div class="footer">
            <p>Powered by Document LLM Agent | AI 驅動的數據分析平台</p>
        </div>
    </div>
</body>
</html>
"""
    return html_template


def generate_markdown_report(title: str, content: str) -> str:
    """Generate a markdown report."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    
    md_template = f"""# 📊 {title}

> 生成時間：{timestamp}
> 類型：AI 專業分析報告

---

{content}

---

*Powered by Document LLM Agent*
"""
    return md_template


def save_report(content: str, filename: str, folder: str = "reports") -> str:
    """Save report to file and return the path."""
    if not os.path.exists(folder):
        os.makedirs(folder)
    
    filepath = os.path.join(folder, filename)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return filepath
